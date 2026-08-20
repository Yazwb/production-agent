# 文件路径 
import logging
from pathlib import Path

from app.domain.models import ParsedPage


logger = logging.getLogger(__name__)


class DocumentParser:
    """将不同格式统一转化为带页码的ParsedPage"""
    SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
    #pages = DocumentParser().parse(path)

    def parse(self, path: Path) -> list[ParsedPage]:
        #将页码进行保存，便于后续召回位置
        suffix = path.suffix.lower()
        logger.info("Parsing document: %s", path.name)
        if suffix not in self.SUPPORTED_SUFFIXES:
            logger.warning("Unsupported document type: %s", suffix or "<none>")
            raise ValueError(f"Unsupport file type: {suffix}")
        if suffix == ".pdf":
            pages = self._parse_pdf(path)
            logger.info("Parsed document %s: %d pages", path.name, len(pages))
            return pages
        if suffix == ".docx":
            pages = self._parse_docx(path)
            logger.info("Parsed document %s: %d pages", path.name, len(pages))
            return pages
        pages = [
            ParsedPage(1, path.read_text(encoding="utf-8", errors="ignore"))
        ]
        logger.info("Parsed document %s: %d pages", path.name, len(pages))
        return pages
        
    @staticmethod
    def _parse_pdf(path: Path) -> list[ParsedPage]:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return [
            ParsedPage(index + 1, page.extract_text() or "")
            for index, page in enumerate(reader.pages)
        ]

    @staticmethod
    def _parse_docx(path: Path) -> list[ParsedPage]:
        from docx import Document

        document = Document(str(path))
        parts = [
            ParsedPage.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        # 一行表格将转化成类似的设备名称 | 型号 | 状态
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip()for cell in row.cells))
        return [
            ParsedPage(1, "\n".join(parts))
        ]
